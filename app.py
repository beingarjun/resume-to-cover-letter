#!/usr/bin/env python3
"""
Resume to Cover Letter Generator
A free Streamlit app using open-source AI models to generate tailored cover letters.
"""

import streamlit as st
import os
import re
import io
from typing import Optional, List, Tuple, Dict, Any
from dataclasses import dataclass
from pathlib import Path

# Document processing imports
try:
    import pypdf
    import docx2txt
    from docx import Document
    from docx.shared import Inches
except ImportError as e:
    st.error(f"Missing dependency: {e}. Run: pip install -r requirements.txt")
    st.stop()

# API clients
import requests
from dotenv import load_dotenv

# Text similarity
try:
    from rapidfuzz import fuzz
except ImportError:
    st.error("Missing rapidfuzz. Run: pip install rapidfuzz")
    st.stop()

# Load environment variables
load_dotenv()

# --- Configuration ---
st.set_page_config(
    page_title="Resume → Cover Letter Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

@dataclass
class APIConfig:
    """Configuration for different AI API providers"""
    name: str
    endpoint: str
    headers: Dict[str, str]
    model: str
    free_tier: bool
    setup_instructions: str

# Available free API configurations
FREE_APIS = {
    "huggingface": APIConfig(
        name="Hugging Face (Free)",
        endpoint="https://api-inference.huggingface.co/models/microsoft/DialoGPT-large",
        headers={},
        model="microsoft/DialoGPT-large",
        free_tier=True,
        setup_instructions="1. Sign up at huggingface.co\n2. Create an access token\n3. Set HUGGINGFACE_API_TOKEN environment variable"
    ),
    "local": APIConfig(
        name="Local Template (No API)",
        endpoint="",
        headers={},
        model="template",
        free_tier=True,
        setup_instructions="No setup required - uses built-in templates"
    )
}

# --- Helper Functions ---

def get_api_client():
    """Get the configured API client"""
    api_choice = st.session_state.get('api_choice', 'local')
    
    if api_choice == 'huggingface':
        token = os.getenv('HUGGINGFACE_API_TOKEN')
        if token:
            return {
                'type': 'huggingface',
                'token': token,
                'endpoint': 'https://api-inference.huggingface.co/models/microsoft/DialoGPT-large'
            }
        else:
            st.warning("Hugging Face token not found. Using local mode.")
            return {'type': 'local'}
    
    return {'type': 'local'}

def extract_text_from_pdf(pdf_file) -> str:
    """Extract text from PDF file"""
    try:
        reader = pypdf.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(docx_file) -> str:
    """Extract text from DOCX file"""
    try:
        return docx2txt.process(docx_file).strip()
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return ""

def extract_text_from_txt(txt_file) -> str:
    """Extract text from TXT file"""
    try:
        return txt_file.read().decode('utf-8').strip()
    except Exception as e:
        st.error(f"Error reading TXT: {e}")
        return ""

def process_uploaded_file(uploaded_file) -> str:
    """Process uploaded resume file and return text"""
    if not uploaded_file:
        return ""
    
    file_type = uploaded_file.type
    
    if file_type == "application/pdf":
        return extract_text_from_pdf(uploaded_file)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(uploaded_file)
    elif file_type == "text/plain":
        return extract_text_from_txt(uploaded_file)
    else:
        st.error(f"Unsupported file type: {file_type}")
        return ""

def analyze_keywords(resume_text: str, jd_text: str) -> Tuple[List[str], List[str]]:
    """Analyze keyword matching between resume and job description"""
    # Extract keywords (simplified approach)
    jd_words = set(re.findall(r'\b[A-Za-z]{3,}\b', jd_text.lower()))
    resume_words = set(re.findall(r'\b[A-Za-z]{3,}\b', resume_text.lower()))
    
    # Common technical and professional keywords
    important_keywords = {
        'python', 'javascript', 'react', 'node', 'sql', 'aws', 'docker', 
        'kubernetes', 'agile', 'scrum', 'leadership', 'management', 'analysis',
        'development', 'design', 'testing', 'api', 'database', 'cloud'
    }
    
    jd_important = jd_words.intersection(important_keywords)
    matched = resume_words.intersection(jd_important)
    missing = jd_important - resume_words
    
    return list(matched), list(missing)

def validate_inputs(resume_text: str, jd_text: str) -> List[str]:
    """Validate user inputs and return list of issues"""
    issues = []
    
    if len(resume_text) < 100:
        issues.append("Resume seems too short. Ensure it includes your full experience.")
    
    if len(jd_text) < 50:
        issues.append("Job description is very short. Include the full job posting for better results.")
    
    if not any(keyword in resume_text.lower() for keyword in ['experience', 'work', 'education', 'skills']):
        issues.append("Resume may be missing key sections (experience, education, skills).")
    
    if not any(keyword in jd_text.lower() for keyword in ['responsibilities', 'requirements', 'experience', 'skills']):
        issues.append("Job description may be incomplete. Include requirements and responsibilities.")
    
    return issues

def generate_cover_letter_local(resume_text: str, jd_text: str, company: str, role: str, 
                               hiring_manager: str, tone: str, length_target: str) -> str:
    """Generate cover letter using local templates (no API required)"""
    
    # Extract key skills and experience from resume
    skills_pattern = r'(?:skills?|technologies?|tools?)[\s\S]*?(?:\n\s*\n|$)'
    skills_match = re.search(skills_pattern, resume_text, re.IGNORECASE)
    skills = skills_match.group(0) if skills_match else "relevant technical skills"
    
    # Extract recent experience
    exp_lines = [line.strip() for line in resume_text.split('\n') if line.strip()]
    recent_exp = ' '.join(exp_lines[:5]) if exp_lines else "my professional experience"
    
    # Generate greeting
    greeting = f"Dear {hiring_manager}," if hiring_manager else "Dear Hiring Manager,"
    
    # Generate opening based on tone
    if tone == "Professional":
        opening = f"I am writing to express my strong interest in the {role} position at {company}. With my background in {skills.lower()}, I am confident I would be a valuable addition to your team."
    elif tone == "Enthusiastic":
        opening = f"I am thrilled to apply for the {role} position at {company}! Your job posting immediately caught my attention, and I am excited about the opportunity to contribute my expertise in {skills.lower()} to your innovative team."
    else:  # Conversational
        opening = f"I came across the {role} opening at {company} and knew I had to apply. My experience with {skills.lower()} aligns perfectly with what you're looking for."
    
    # Generate body paragraphs
    body1 = f"In my recent roles, {recent_exp[:200]}... I have developed strong capabilities that directly match your requirements."
    
    body2 = "I am particularly drawn to this opportunity because it combines my technical expertise with my passion for innovation. I would welcome the chance to discuss how my background can contribute to your team's success."
    
    # Generate closing
    closing = f"Thank you for considering my application. I look forward to hearing from you soon.\n\nSincerely,\n[Your Name]"
    
    # Combine all parts
    cover_letter = f"{greeting}\n\n{opening}\n\n{body1}\n\n{body2}\n\n{closing}"
    
    return cover_letter

def generate_cover_letter_api(client: dict, resume_text: str, jd_text: str, 
                             company: str, role: str, hiring_manager: str, 
                             tone: str, length_target: str) -> str:
    """Generate cover letter using API"""
    
    if client['type'] == 'huggingface':
        # Simplified prompt for Hugging Face
        prompt = f"Write a professional cover letter for {role} at {company} based on this resume and job description. Tone: {tone}"
        
        try:
            response = requests.post(
                client['endpoint'],
                headers={"Authorization": f"Bearer {client['token']}"},
                json={"inputs": prompt[:512]}  # Limit input length
            )
            
            if response.status_code == 200:
                result = response.json()
                if isinstance(result, list) and len(result) > 0:
                    return result[0].get('generated_text', '')
                else:
                    st.warning("API returned unexpected format. Using local generation.")
                    return generate_cover_letter_local(resume_text, jd_text, company, role, hiring_manager, tone, length_target)
            else:
                st.warning(f"API request failed: {response.status_code}. Using local generation.")
                return generate_cover_letter_local(resume_text, jd_text, company, role, hiring_manager, tone, length_target)
                
        except Exception as e:
            st.warning(f"API error: {e}. Using local generation.")
            return generate_cover_letter_local(resume_text, jd_text, company, role, hiring_manager, tone, length_target)
    
    # Fallback to local generation
    return generate_cover_letter_local(resume_text, jd_text, company, role, hiring_manager, tone, length_target)

def to_docx(text: str) -> bytes:
    """Convert text to DOCX format"""
    doc = Document()
    
    # Add paragraphs
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')
    
    # Save to bytes
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Main Application ---

def main():
    # Header
    st.title("📄 Resume → Cover Letter Generator")
    st.markdown("**Transform your resume into tailored cover letters using free AI models**")
    
    # Sidebar for API configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        api_choice = st.selectbox(
            "Choose AI Provider:",
            options=['local', 'huggingface'],
            format_func=lambda x: FREE_APIS[x].name,
            key='api_choice'
        )
        
        if api_choice != 'local':
            st.info(FREE_APIS[api_choice].setup_instructions)
            
            if api_choice == 'huggingface':
                token_status = "✅ Token found" if os.getenv('HUGGINGFACE_API_TOKEN') else "❌ Token missing"
                st.write(f"Status: {token_status}")
        
        st.header("📝 Writing Options")
        tone = st.selectbox("Tone:", ["Professional", "Enthusiastic", "Conversational"])
        length_target = st.selectbox("Length:", ["Concise (3-4 paragraphs)", "Standard (4-5 paragraphs)", "Detailed (5-6 paragraphs)"])
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📄 Upload Your Resume")
        
        # Enhanced file upload guidance
        with st.expander("📋 Resume Upload Tips", expanded=False):
            st.markdown("""
            **For best results:**
            - Use a well-formatted resume with clear sections
            - Include your full work history, skills, and education
            - Ensure text is selectable (not just images)
            - Supported formats: PDF, DOCX, TXT
            
            **What makes a good resume:**
            - Clear contact information
            - Professional summary or objective
            - Detailed work experience with achievements
            - Relevant skills and technologies
            - Education and certifications
            """)
        
        uploaded_file = st.file_uploader(
            "Choose your resume file",
            type=['pdf', 'docx', 'txt'],
            help="Upload PDF, DOCX, or TXT format. Make sure text is selectable for best results."
        )
        
        resume_text = ""
        if uploaded_file:
            with st.spinner("Processing resume..."):
                resume_text = process_uploaded_file(uploaded_file)
            
            if resume_text:
                st.success(f"✅ Resume loaded ({len(resume_text)} characters)")
                with st.expander("Preview resume content", expanded=False):
                    st.text_area("Resume text:", resume_text[:500] + "..." if len(resume_text) > 500 else resume_text, height=150, disabled=True)
            else:
                st.error("❌ Could not extract text from file. Please check the format.")
    
    with col2:
        st.subheader("💼 Job Description")
        
        # Enhanced job description guidance
        with st.expander("📋 Job Description Tips", expanded=False):
            st.markdown("""
            **For best results:**
            - Paste the complete job posting
            - Include requirements, responsibilities, and qualifications
            - Don't edit or summarize - use the full text
            - Include company information if available
            
            **Good job descriptions include:**
            - Role title and department
            - Key responsibilities
            - Required skills and experience
            - Preferred qualifications
            - Company culture information
            """)
        
        jd_text = st.text_area(
            "Paste the complete job description here:",
            height=200,
            placeholder="Copy and paste the full job posting, including requirements, responsibilities, and company information...",
            help="Include the complete job posting for better keyword matching and personalization."
        )
        
        if jd_text:
            st.success(f"✅ Job description loaded ({len(jd_text)} characters)")
    
    # Additional details section
    st.subheader("🎯 Additional Details (Optional)")
    
    col3, col4, col5 = st.columns(3)
    with col3:
        company = st.text_input("Company Name:", placeholder="e.g., TechCorp Inc.")
    with col4:
        role = st.text_input("Job Title:", placeholder="e.g., Senior Software Engineer")
    with col5:
        hiring_manager = st.text_input("Hiring Manager:", placeholder="e.g., Sarah Johnson")
    
    # Analysis section
    if resume_text and jd_text:
        st.subheader("🔍 Smart Analysis")
        
        # Validate inputs
        issues = validate_inputs(resume_text, jd_text)
        matched_keywords, missing_keywords = analyze_keywords(resume_text, jd_text)
        
        col_analysis1, col_analysis2 = st.columns(2)
        
        with col_analysis1:
            st.write("**✅ Matched Keywords:**")
            if matched_keywords:
                st.write(", ".join(matched_keywords[:10]))  # Show first 10
            else:
                st.write("(few matches detected - consider adding relevant skills)")
        
        with col_analysis2:
            st.write("**⚠️ Missing Keywords:**")
            if missing_keywords:
                st.write(", ".join(missing_keywords[:10]))  # Show first 10
                st.caption("Consider incorporating these terms naturally in your cover letter.")
            else:
                st.write("(good keyword coverage detected)")
        
        if issues:
            st.warning("**Suggestions for improvement:**\n" + "\n".join(f"• {issue}" for issue in issues))
    
    # Generation section
    st.divider()
    
    generate_btn = st.button(
        "✨ Generate Cover Letter",
        type="primary",
        disabled=not (resume_text and jd_text),
        help="Generate a tailored cover letter based on your resume and the job description"
    )
    
    # Store generated text in session state
    if "generated_text" not in st.session_state:
        st.session_state.generated_text = ""
    
    if generate_btn:
        if not resume_text:
            st.warning("⚠️ Please upload your resume first.")
        elif not jd_text:
            st.warning("⚠️ Please paste the job description.")
        else:
            client = get_api_client()
            
            with st.spinner("🤖 Crafting your personalized cover letter..."):
                try:
                    generated_letter = generate_cover_letter_api(
                        client=client,
                        resume_text=resume_text,
                        jd_text=jd_text,
                        company=company,
                        role=role,
                        hiring_manager=hiring_manager,
                        tone=tone,
                        length_target=length_target
                    )
                    st.session_state.generated_text = generated_letter
                except Exception as e:
                    st.error(f"Generation failed: {e}")
                    st.session_state.generated_text = ""
    
    # Display and edit generated content
    if st.session_state.generated_text:
        st.subheader("📝 Your Generated Cover Letter")
        
        # Editable text area
        edited_text = st.text_area(
            "Review and edit your cover letter:",
            value=st.session_state.generated_text,
            height=400,
            help="Make any changes you'd like before downloading or copying."
        )
        
        # Update session state with edits
        st.session_state.generated_text = edited_text
        
        # Export options
        col_export1, col_export2, col_export3 = st.columns(3)
        
        with col_export1:
            # Download as DOCX
            filename = f"Cover_Letter_{(company or 'Application').replace(' ', '_')}.docx"
            st.download_button(
                "⬇️ Download as DOCX",
                data=to_docx(edited_text),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download as a Microsoft Word document"
            )
        
        with col_export2:
            # Copy to clipboard
            if st.button("📋 Copy to Clipboard"):
                st.code(edited_text, language=None)
                st.success("✅ Cover letter copied! Paste it wherever you need it.")
        
        with col_export3:
            # Clear/start over
            if st.button("🔄 Start Over"):
                st.session_state.generated_text = ""
                st.rerun()
    
    # Footer with tips
    st.divider()
    
    with st.expander("💡 Pro Tips for Great Cover Letters", expanded=False):
        st.markdown("""
        **Before submitting:**
        - ✅ Proofread for spelling and grammar
        - ✅ Customize the greeting with the hiring manager's name if known
        - ✅ Add specific examples from your experience
        - ✅ Match the company's tone and culture
        - ✅ Keep it to one page
        
        **What employers want to see:**
        - Clear connection between your skills and their needs
        - Specific examples of relevant achievements
        - Enthusiasm for the role and company
        - Professional formatting and language
        - Evidence that you researched the company
        """)

if __name__ == "__main__":
    main()